from ragdocforge.parsers.base_parser import base_document
from ragdocforge.schemas.document_models import ParsedDocument


class DocxParser:
    def parse(self, path: str) -> ParsedDocument:
        try:
            import docx
        except ImportError:
            return base_document(path, "", ["DOCX parsing requires python-docx."])

        parts: list[str] = []
        warnings: list[str] = []
        try:
            document = docx.Document(path)
            for paragraph in document.paragraphs:
                text = paragraph.text.strip()
                if not text:
                    continue
                style = paragraph.style.name.lower() if paragraph.style else ""
                parts.append(f"## {text}" if "heading" in style else text)
            for table in document.tables:
                for row in table.rows:
                    cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                    parts.append("| " + " | ".join(cells) + " |")
        except Exception as exc:  # pragma: no cover - depends on external files
            warnings.append(f"DOCX extraction failed: {exc}")
        return base_document(path, "\n".join(parts), warnings)
